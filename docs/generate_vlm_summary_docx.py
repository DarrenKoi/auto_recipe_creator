"""VLM 설치 과정 및 벤치마크 요약 문서 생성 스크립트"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path


def set_cell_text(cell, text, bold=False, size=9):
    cell.text = ""
    run = cell.paragraphs[0].add_run(text)
    run.font.size = Pt(size)
    run.bold = bold


def build_document():
    doc = Document()

    # --- 스타일 설정 ---
    style = doc.styles["Normal"]
    style.font.name = "맑은 고딕"
    style.font.size = Pt(10)

    # === 제목 ===
    title = doc.add_heading("VLM 설치 과정 및 ScreenSpot-Pro 벤치마크 요약", level=1)
    title.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

    # === 1. VLM 설치 과정 및 목적 ===
    doc.add_heading("1. VLM 설치 과정 및 목적", level=2)
    doc.add_paragraph(
        "HuggingFace에서 모델 가중치를 다운로드한 뒤, 인터넷이 차단된 자사 Private Cloud 서버로 "
        "수동 이관하여 vLLM 기반으로 서빙했다. H200 140GB GPU 2장에 총 5개 VLM을 "
        "co-location 방식으로 분배 배치했다."
    )

    doc.add_heading("GPU 할당", level=3)
    gpu_items = [
        "GPU 0: UI-Venus-1.5-8B (포트 8001, Primary) + UI-TARS-1.5-7B (포트 8003, 대안 비교용) — auto-tune 메모리 활용",
        "GPU 1: MAI-UI-8B (포트 8002, crop 재분석) + PaddleOCR-VL-1.5 (포트 8004, OCR 0.9B) + GOT-OCR-2.0 (포트 8005, OCR fallback)",
    ]
    for item in gpu_items:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("각 모델의 역할", level=3)
    roles = [
        ("UI-Venus-1.5-8B", "전체 화면 GUI 요소 좌표 탐지의 주력 모델 (ScreenSpot-Pro 69.6%, SOTA)"),
        ("UI-TARS-1.5-7B", "Qwen2.5-VL 기반의 에이전트형 grounding 비교 모델"),
        ("MAI-UI-8B", "밀집 UI 영역을 crop하여 2차 정밀 분석하는 보조 모델"),
        ("PaddleOCR-VL-1.5", "레이아웃 이해 기반 OCR 엔진 (0.9B 경량)"),
        ("GOT-OCR-2.0", "텍스트 재확인용 fallback, Transformers 직접 추론 방식"),
    ]
    for name, desc in roles:
        p = doc.add_paragraph(style="List Bullet")
        run_bold = p.add_run(f"{name}: ")
        run_bold.bold = True
        p.add_run(desc)

    doc.add_heading("인프라 구성", level=3)
    doc.add_paragraph(
        "Flask 프록시가 미들웨어로 동작하여 사내 Windows PC에서 service slug 기반 URL "
        "(/api/vlm_serve/{slug}/v1/chat/completions)로 통합 라우팅한다. "
        "모든 모델은 오프라인 정책을 준수하며, 로컬 절대경로로만 로딩된다."
    )

    doc.add_heading("최종 목표", level=3)
    doc.add_paragraph(
        "CD-SEM/VeritySEM 장비의 RCS 소프트웨어 화면을 VLM으로 분석하고, "
        "GUI 자동화(pywinauto/pynput)로 레시피 생성 과정을 자동화하여 "
        "반도체 계측 장비의 수동 레시피 셋업을 대체하는 것이다."
    )

    # === 2. ScreenSpot-Pro 벤치마크 비교 ===
    doc.add_heading("2. ScreenSpot-Pro 벤치마크 성능 비교", level=2)

    # 테이블
    table_data = [
        ("모델", "파라미터", "ScreenSpot-Pro", "비고"),
        ("UI-Venus-1.5", "8B", "69.6%", "현재 SOTA, 사내 Primary"),
        ("Qwen3-VL", "32B", "61.8%", "사내 운영 중"),
        ("Kimi-VL (K2.5 계열)", "MoE", "52.8%", "사내 운영 중"),
        ("UI-TARS-1.5", "7B", "35.7%", "사내 비교용"),
        ("OS-Atlas", "7B", "18.9%", "이전 세대"),
        ("GPT-4o", "-", "<2%", "범용 MLLM"),
    ]

    table = doc.add_table(rows=len(table_data), cols=4)
    table.style = "Light Shading Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row_data in enumerate(table_data):
        for col_idx, cell_text in enumerate(row_data):
            is_header = row_idx == 0
            set_cell_text(table.rows[row_idx].cells[col_idx], cell_text, bold=is_header)

    doc.add_paragraph("")  # spacing
    doc.add_paragraph(
        "UI-Venus가 8B 모델임에도 불구하고 Qwen3-VL(32B) 대비 +7.8%p, "
        "Kimi-VL 대비 +16.8%p 높은 정확도를 보인다. 파라미터 수가 4배 적으면서도 "
        "GUI grounding에 특화되어 더 높은 점수를 달성한 것이 핵심이며, "
        "이것이 Primary 모델로 UI-Venus를 선택한 이유이다."
    )

    # === 3. 왜 %로 벤치마크를 표시하는가 ===
    doc.add_heading("3. 왜 %로 벤치마크를 표시하는가", level=2)
    doc.add_paragraph(
        "ScreenSpot-Pro는 1,581개의 전문가 주석 스크린샷-지시문 쌍으로 구성되어 있다. "
        "평가 방식은 point-in-box accuracy로, 모델이 예측한 좌표가 정답 바운딩 박스 안에 "
        "들어가면 정답, 아니면 오답이다."
    )
    p = doc.add_paragraph()
    run = p.add_run("정확도(%) = (정답 수 / 전체 1,581문제) × 100")
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    reasons = [
        "데이터셋 크기가 다른 벤치마크와 비교 가능 — 1,581개든 5,000개든 비율로 정규화하면 직접 비교 가능",
        "직관적 해석 — \"69.6%\"는 \"100번 중 약 70번 정확히 클릭한다\"로 바로 이해됨",
        "이진 판정(맞다/틀리다)의 집계 — GUI grounding은 \"박스 안에 들어갔는가\"의 Yes/No 판정이므로, "
        "이를 집계하면 자연스럽게 비율(%)이 된다",
    ]
    for reason in reasons:
        doc.add_paragraph(reason, style="List Number")

    doc.add_paragraph(
        "절대 점수(예: 1,099/1,581)로 표기하면 데이터셋 규모를 모르는 사람은 해석이 불가능하기 때문에, "
        "모든 ML 벤치마크가 %를 표준으로 사용한다."
    )

    # === 출처 ===
    doc.add_heading("출처", level=2)
    sources = [
        "ScreenSpot-Pro Leaderboard — https://gui-agent.github.io/grounding-leaderboard/",
        "UI-Venus GitHub — https://github.com/inclusionAI/UI-Venus",
        "ScreenSpot-Pro 논문 (arXiv) — https://arxiv.org/abs/2504.07981",
        "Qwen3-VL 벤치마크 — https://the-decoder.com/qwen3-vl-can-scan-two-hour-videos-and-pinpoint-nearly-every-detail/",
        "Kimi-VL GitHub — https://github.com/MoonshotAI/Kimi-VL",
        "Kimi K2.5 Tech Blog — https://www.kimi.com/blog/kimi-k2-5",
    ]
    for src in sources:
        doc.add_paragraph(src, style="List Bullet")

    # 저장
    output_path = Path(__file__).parent / "vlm_summary_benchmark.docx"
    doc.save(str(output_path))
    print(f"[INFO] 문서 생성 완료: {output_path}")


if __name__ == "__main__":
    build_document()
