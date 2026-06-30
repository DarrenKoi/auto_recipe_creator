"""프로젝트 진행 보고서 Word(.docx) 생성 스크립트.

docs/project_progress/ 의 Markdown 보고서(00~05)를 **그대로 읽어** 하나의 임원 보고용 .docx 로
렌더링한다. 보고서 내용 수정은 .md 만 고치면 되며(단일 source-of-truth), 본 스크립트는
스타일·레이아웃만 담당한다.

표지 -> 목차 -> 00~05 본문(파일별 페이지 구분) 순. 시각 보조 자료는 짝 문서 `_appendix.html` 참조.

실행:
  uv run python docs/project_progress/build_report_docx.py
출력:
  docs/project_progress/project_progress_report.docx

본문은 한국어 경어체, 모델명/CV 기법/env 등 기술 용어는 영문 병기. 한글 폰트는 맑은 고딕.
"""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from _md_report import load_report_docs, parse_inline

OUTPUT_PATH = Path(__file__).resolve().parent / "project_progress_report.docx"
BASE_DIR = Path(__file__).resolve().parent

KR_FONT = "맑은 고딕"
MONO_FONT = "Consolas"

NAVY = RGBColor(0x14, 0x2A, 0x55)
ACCENT = RGBColor(0xE8, 0x6A, 0x1F)
MUTED = RGBColor(0x55, 0x60, 0x70)
CODE_BG = RGBColor(0x33, 0x3A, 0x45)


def _apply_font(run, *, name=KR_FONT, size=None, bold=None, color=None):
    """run 에 폰트를 적용한다. 한글은 w:eastAsia 속성까지 설정해야 제대로 렌더된다."""
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = rpr.makeelement(qn("w:rFonts"), {})
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), name)
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    if color is not None:
        run.font.color.rgb = color


def _add_inline(paragraph, text, *, size=10.5, base_color=None):
    """**bold**/`code`/[link] 인라인 서식을 해석해 run 들을 추가한다."""
    for seg in parse_inline(text):
        run = paragraph.add_run(seg["text"])
        if seg["mono"]:
            _apply_font(run, name=MONO_FONT, size=size - 0.5, color=MUTED)
        else:
            _apply_font(run, size=size, bold=seg["bold"] or None, color=base_color)


def _para(doc, text="", *, size=10.5, bold=False, color=None, align=None,
          space_after=6, space_before=0, mono=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    if align is not None:
        p.alignment = align
    if text:
        run = p.add_run(text)
        _apply_font(run, name=MONO_FONT if mono else KR_FONT, size=size,
                    bold=bold, color=color)
    return p


def _rich_para(doc, text, *, size=10.5, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    _add_inline(p, text, size=size)
    return p


def _heading(doc, text, *, level=1):
    sizes = {1: 17, 2: 13.5, 3: 11.5}
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(6)
    # heading 텍스트는 단일 스타일(인라인 서식 제거)
    plain = "".join(seg["text"] for seg in parse_inline(text))
    run = p.add_run(plain)
    _apply_font(run, size=sizes.get(level, 12), bold=True,
                color=NAVY if level <= 2 else ACCENT)
    return p


def _list_item(doc, item, *, ordered):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(18 + item["level"] * 16)
    p.paragraph_format.space_after = Pt(3)
    if ordered:
        marker, mcolor = item["marker"] + " ", NAVY
    else:
        marker, mcolor = ("• " if item["level"] == 0 else "◦ "), ACCENT
    mrun = p.add_run(marker)
    _apply_font(mrun, size=10.5, bold=True, color=mcolor)
    _add_inline(p, item["text"], size=10.5)
    return p


def _callout(doc, text):
    """blockquote -> 좌측 강조선 느낌의 콜아웃(들여쓰기 + muted)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.space_before = Pt(4)
    bar = p.add_run("▎ ")
    _apply_font(bar, size=10.5, bold=True, color=ACCENT)
    _add_inline(p, text, size=10, base_color=MUTED)
    return p


def _table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        _add_inline_cell(hdr[i].paragraphs[0], h, size=9.5, bold=True,
                         color=RGBColor(0xFF, 0xFF, 0xFF))
    for row in rows:
        cells = table.add_row().cells
        for i in range(len(headers)):
            val = row[i] if i < len(row) else ""
            cells[i].text = ""
            _add_inline_cell(cells[i].paragraphs[0], val, size=9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def _add_inline_cell(paragraph, text, *, size=9.5, bold=False, color=None):
    for seg in parse_inline(text):
        run = paragraph.add_run(seg["text"])
        if seg["mono"]:
            _apply_font(run, name=MONO_FONT, size=size - 0.5,
                        bold=bold or None, color=color)
        else:
            _apply_font(run, size=size, bold=(bold or seg["bold"]) or None, color=color)


def _set_cell_shading(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.makeelement(qn("w:shd"), {})
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _shade_header(table, hex_color="142A55"):
    for cell in table.rows[0].cells:
        _set_cell_shading(cell, hex_color)


def _code_block(doc, lines):
    for line in lines:
        _para(doc, line if line else " ", size=9, mono=True, color=MUTED, space_after=0)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


# ---------------------------------------------------------------------------
# 표지 / 목차 / 본문 렌더링
# ---------------------------------------------------------------------------


def build_cover(doc):
    for _ in range(4):
        doc.add_paragraph()
    _para(doc, "프로젝트 진행 보고서", size=30, bold=True, color=NAVY,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    _para(doc, "AI 기반 CD-SEM / VeritySEM Recipe 자동 Setup PoC", size=15,
          color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    _para(doc, "VLM 배포·운영  ·  workflow_1 / workflow_2 / workflow_3", size=12,
          color=ACCENT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=40)
    _para(doc, "목적 · PoC 방향 · 성과 · 확장성", size=11, color=MUTED,
          align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6)
    _para(doc, "(시각 보조 자료: 짝 문서 _appendix.html 참조)", size=9.5, color=MUTED,
          align=WD_ALIGN_PARAGRAPH.CENTER)
    doc.add_page_break()


def build_toc(doc, docs):
    _heading(doc, "목차 (Contents)", level=1)
    for d in docs:
        title = next((b["text"] for b in d["blocks"] if b["type"] == "h" and b["level"] == 1),
                     d["name"])
        plain = "".join(seg["text"] for seg in parse_inline(title))
        _para(doc, plain, size=11, space_after=4)
    doc.add_page_break()


def render_blocks(doc, blocks):
    for b in blocks:
        t = b["type"]
        if t == "h":
            _heading(doc, b["text"], level=min(max(b["level"], 1), 3))
        elif t == "p":
            _rich_para(doc, b["text"])
        elif t == "quote":
            _callout(doc, b["text"])
        elif t == "code":
            _code_block(doc, b["lines"])
        elif t == "list":
            for item in b["items"]:
                _list_item(doc, item, ordered=b["ordered"])
        elif t == "table":
            tbl = _table(doc, b["headers"], b["rows"])
            _shade_header(tbl)


def main():
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = KR_FONT
    normal.font.size = Pt(10.5)
    normal_rpr = normal.element.get_or_add_rPr()
    rfonts = normal_rpr.makeelement(qn("w:rFonts"), {})
    rfonts.set(qn("w:eastAsia"), KR_FONT)
    normal_rpr.append(rfonts)

    docs = load_report_docs(BASE_DIR)

    build_cover(doc)
    build_toc(doc, docs)
    for idx, d in enumerate(docs):
        if idx > 0:
            doc.add_page_break()
        render_blocks(doc, d["blocks"])

    doc.save(OUTPUT_PATH)
    print(f"[INFO] DOCX saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
